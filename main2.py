from UI.main_window import Ui_MainWindow #главное окно
from UI.menu_options import Ui_dialog_options # меню настройки
from UI.antenna_options import Ui_dialog_antenns # меню антенн
from rec_func import main_copm #модуль расчета
from PyQt5 import QtWidgets, QtCore, QtGui
import sys, os
import numpy as np
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

#главное_окно
class Main_Window(QtWidgets.QMainWindow):
    def __init__(self):
        QtWidgets.QWidget.__init__(self)
        self.ui = Ui_MainWindow()
        self.ui.setupUi(self)

        #виджеты главного окна
        self.ui.pushButton_1.clicked.connect(self.add_frequence)
        self.ui.treeWidget.itemClicked.connect(self.item_on_clicked)
        self.ui.pushButton_2.clicked.connect(self.delete_frequence)
        self.ui.pushButton_3.clicked.connect(self.choice_file)
        self.ui.pushButton_4.clicked.connect(self.save_frequence_file)
        self.ui.pushButton_5.clicked.connect(self.make_rec)# для расчета
        self.ui.pushButton_6.clicked.connect(self.open_project)
        self.ui.pushButton_7.clicked.connect(self.save_project)
        self.ui.menu_options.triggered.connect(self.options_window)
        self.ui.menu_antena.triggered.connect(self.options_antenns)
        self.ui.eksport_word.triggered.connect(self.eksport_word_func)
        #изменение объекта древа
        self.ui.treeWidget.itemChanged.connect(self.rename_item_on_itemChanged)

    #обработка изменения переименования частоты
    #добавляет файл .txt и удаляет старый .txt
    def rename_item_on_itemChanged(self, item):
        #проверка на выделенный объект
        # try:
        items = self.ui.treeWidget.selectedItems()
        if float(item.text(0)) < 100:
            error = QtWidgets.QMessageBox(QtWidgets.QMessageBox.Information, 'Ошибка', 'Частота не может быть меньше 100, задайте другую частоту: rename_item_on_itemChanged',
                                          QtWidgets.QMessageBox.Ok)
            error.exec_()
        name = float(item.text(0))
        # save_folder = str(sys.argv)[2:str(sys.argv).rfind('/') + 1] + 'Files/' + item.text(0) + '.txt'
        save_folder = str(sys.argv)[2:str(sys.argv).rfind('/') + 1] + 'Files/' + str(name) + '.txt'
        arr = np.array([])

        if items:
            item_first = self.ui.tableWidget.item(0, 0)
            if item_first != None:
                # создаем новый файл с новой частотой
                for row in range(0, self.ui.tableWidget.rowCount()):
                    for column in range(0, self.ui.tableWidget.columnCount()):
                        item = self.ui.tableWidget.item(row, column)
                        arr = np.append(arr, item.text())
                arr = np.array(arr, dtype=np.float64)
                arr = np.reshape(arr, [self.ui.tableWidget.rowCount(), 3])
                with open(save_folder, 'w') as file:
                    for line in arr:
                        for x in line:
                            print(x, file=file, end='\t')
                        print('\n', file=file, end='')
                print('file created success')
                # удаляем прошлый файл с частотой
                mas_f = os.listdir(save_folder[:save_folder.rfind('/') + 1])
                mas_f_real = np.array([])

                # считываем по-элементно

                for x in range(0, self.ui.treeWidget.topLevelItemCount()):
                    item = self.ui.treeWidget.topLevelItem(x)
                    mas_f_real = np.append(mas_f_real, str(float(item.text(0))))

                # добавляем .txt
                for x in range(0, len(mas_f_real)):
                    mas_f_real[x] = mas_f_real[x] + '.txt'

                # вычисляем лишний файл
                file_name = str(set.difference(set(mas_f), set(mas_f_real)))[2:-2]
                try:
                    os.remove(str(sys.argv)[2:str(sys.argv).rfind('/') + 1] + 'Files/' + file_name)
                except:
                    pass

            else:
                # создаем новый пустой файл .txt
                with open(save_folder, 'w') as file:
                    file.close()
                # удаляем прошлый файл с частотой
                mas_f = os.listdir(save_folder[:save_folder.rfind('/') + 1])
                mas_f_real = np.array([])
                # считываем по-элементно
                for x in range(0, self.ui.treeWidget.topLevelItemCount()):
                    item = self.ui.treeWidget.topLevelItem(x)
                    mas_f_real = np.append(mas_f_real, str(float(item.text(0))))
                # добавляем .txt
                for x in range(0, len(mas_f_real)):
                    mas_f_real[x] = mas_f_real[x] + '.txt'
                # вычисляем лишний файл
                file_name = str(set.difference(set(mas_f), set(mas_f_real)))[2:-2]
                os.remove(str(sys.argv)[2:str(sys.argv).rfind('/') + 1] + 'Files/' + file_name)

            self.ui.statusbar.showMessage('Частота успешно переименована')

        # except:
        #     error = QtWidgets.QMessageBox(QtWidgets.QMessageBox.Information, 'Ошибка', 'Ошибка переименования частоты : rename_item_on_itemChanged',
        #                                   QtWidgets.QMessageBox.Ok)
        #     error.exec_()



    #инициализация стандартных настроект переменных, запускается один раз при старте программ
    def init_settings(self):
        #глобальный доступ к словарю с переменными
        global settings_dict
        keys = []
        values = []
        # считывает переменные из settings.txt
        settings_folder = str(sys.argv)[2:str(sys.argv).rfind('/') + 1] + 'setup/settings.txt'
        # создаем словарь для доступа к данным
        try:
            file = np.loadtxt(settings_folder, delimiter='=', dtype=np.unicode)
            for line in file:
                keys.append(line[0])
                values.append(line[1])
            settings_dict = dict(zip(keys, values))
            #строка состояния
            self.ui.statusbar.showMessage('Настройки загружены')
        # если не удается прочитать словарь...
        except:
            error = QtWidgets.QMessageBox(QtWidgets.QMessageBox.Information, 'Ошибка', 'Ошибка инициализации словаря переменных : init_settings',
                                          QtWidgets.QMessageBox.Ok)
            error.exec_()


    #открывает окно настроект антенн
    def options_antenns(self):
        global antenns_dict #массив с выбранными антеннами для расчета
        self.window = menu_antenns_options_window()
        #инициализация = файл с антеннами
        self.window.init_settings()
        self.window.show()
        # #инициализация данных в listwidget
        settings_folder = str(sys.argv)[2:str(sys.argv).rfind('/') + 1] + 'antens/'
        self.window.listWidget.clear()
        for file in antenn_files:
            self.window.listWidget.addItem(file)
        self.window.listWidget.update()
        #инициализация данных в combo-boxs
        self.window.comboBox.clear()
        self.window.comboBox.addItems(antenn_files)
        self.window.comboBox.update()
        self.window.comboBox_2.addItems(antenn_files)
        self.window.comboBox_2.update()
        self.window.comboBox_3.addItems(antenn_files)
        self.window.comboBox_3.update()
        self.window.comboBox_4.addItems(antenn_files)
        self.window.comboBox_4.update()
        self.window.show()

        result = self.window.exec()
        # print(result)
        if result == 1:
        #инициализация массива работающих антенн
            keys = ['PRM_NH', 'PRM_VH', 'PRD_NH', 'PRD_VH']
            values = []
            values.append(self.window.comboBox.currentText())
            values.append(self.window.comboBox_2.currentText())
            values.append(self.window.comboBox_3.currentText())
            values.append(self.window.comboBox_4.currentText())
            antenns_dict = dict(zip(keys,values))
        #заполняем поля выбранными антеннами в проекте на главном окне
            self.ui.lineEdit_3.setText(str(antenns_dict['PRM_NH']))
            self.ui.lineEdit_3.update()
            self.ui.lineEdit_5.setText(str(antenns_dict['PRM_VH']))
            self.ui.lineEdit_5.update()
            self.ui.lineEdit_4.setText(str(antenns_dict['PRD_NH']))
            self.ui.lineEdit_4.update()
            self.ui.lineEdit_6.setText(str(antenns_dict['PRD_VH']))
            self.ui.lineEdit_6.update()
        #чек бокс ставим на true
            self.ui.checkBox.setChecked(True)
            self.ui.checkBox.update()
            # print(antenns_dict)

    # открывает окно настроек_переменных из меню
    def options_window(self):
        self.window = menu_options_window()
        #путь к файлу настроек settings.txt
        settings_folder = str(sys.argv)[2:str(sys.argv).rfind('/') + 1] + 'setup/settings.txt'
        # print(settings_folder)
        #задaем значения полям
        try:
            self.window.doubleSpinBox.setValue(float(settings_dict['Rprd']))
            self.window.doubleSpinBox_2.setValue(float(settings_dict['Rprm']))
            self.window.doubleSpinBox_3.setValue(float(settings_dict['Gym']))
            self.window.doubleSpinBox_4.setValue(float(settings_dict['Pgvc_AR_stac']))
            self.window.doubleSpinBox_5.setValue(float(settings_dict['Pgvc_AR_port']))

            # настройка combobox-a
            if settings_dict['typeVP'] == 'ВП с СЗУ':
                self.window.comboBox.setCurrentIndex(0)  # выбрать 0ый итем
            elif settings_dict['typeVP'] == 'ВП без СЗУ':
                self.window.comboBox.setCurrentIndex(1)  # выбрать 1ый итем

            self.window.show()

            result = self.window.exec_()
            # print(result)

            if result == 1:  # это Ok
            # считываем новые данные из полей
                settings_dict['Rprd'] = self.window.doubleSpinBox.value()
                settings_dict['Rprm'] = self.window.doubleSpinBox_2.value()
                settings_dict['Gym'] = self.window.doubleSpinBox_3.value()
                settings_dict['Pgvc_AR_stac'] = self.window.doubleSpinBox_4.value()
                settings_dict['Pgvc_AR_port'] = self.window.doubleSpinBox_5.value()
                settings_dict['typeVP'] = self.window.comboBox.currentText()
                with open(settings_folder, 'w') as file:
                    for key in settings_dict.keys():
                        print(key + '=' + str(settings_dict[key]), file=file, end='\n')
                self.ui.statusbar.showMessage('Настройки успешно сохранены в проект')
            elif result != 0:
                print('ошибка сохранения словаря settings')
                return

        # если файл settings будет изменен, ошибка словаря
        except:
            error = QtWidgets.QMessageBox(QtWidgets.QMessageBox.Information, 'Ошибка', 'Ошибка словаря переменных : options_window',
                                          QtWidgets.QMessageBox.Ok)
            error.exec_()

    #добавить частоту
    def add_frequence(self):
        # ok - флаг
        try:
            frequence, ok = QtWidgets.QInputDialog.getDouble(QtWidgets.QWidget(),
                                                             'Окно ввода', 'Введите частоту ВЧО в МГц')
            #проверка на частоту
            if frequence < 100:
                error = QtWidgets.QMessageBox(QtWidgets.QMessageBox.Information, 'Ошибка',
                                              'Вводимая частота не может быть меньше 100 МГц', QtWidgets.QMessageBox.Ok)
                error.exec()
                return

            frequence = str(frequence)

            if not ok:
                return
            tree_row = QtWidgets.QTreeWidgetItem(self.ui.treeWidget)
            tree_row.setText(0, frequence)
            tree_row.setTextAlignment(0,QtCore.Qt.AlignCenter)
            tree_row.setFlags(tree_row.flags() | QtCore.Qt.ItemIsEditable)
            # tree_row.setFlags(tree_row.flags()| QtCore.Qt.ItemIsUserCheckable)
            tree_row.setCheckState(0, QtCore.Qt.Unchecked)
            self.ui.treeWidget.update()

            #создаем файл при создании частоты
            with open(str(sys.argv)[2:str(sys.argv).rfind('/') + 1] + 'Files/' + frequence + '.txt', 'w') as file:
                file.close()
            self.ui.statusbar.showMessage('Частота успешно добавлена')
        except:
            error = QtWidgets.QMessageBox(QtWidgets.QMessageBox.Information, 'Ошибка', 'Ошибка добавления частоты : add_frequence',
                                          QtWidgets.QMessageBox.Ok)
            error.exec_()

    #удалить частоту
    def delete_frequence(self):

        items = self.ui.treeWidget.selectedItems()
        root = self.ui.treeWidget.invisibleRootItem()
        if not items : return
        for item in items:
            root.removeChild(item)
            # print(item.text(0))
        #также удаляет файл txt
            try:
                os.remove(str(sys.argv)[2:str(sys.argv).rfind('/') + 1] + 'Files/' + str(float(item.text(0))) + '.txt')
                self.ui.statusbar.showMessage('Частота успешно удалена')
                # print(str(sys.argv)[2:str(sys.argv).rfind('/') + 1] + 'Files/' + item.text() + '.txt')
            except:
                error = QtWidgets.QMessageBox(QtWidgets.QMessageBox.Information, 'Ошибка',
                                              'Ошибка удаления частоты : delete_frequence',
                                              QtWidgets.QMessageBox.Ok)
                error.exec_()

    #выбрать файл для загрузки данных в таблицу по выбранной частоте
    def choice_file(self):
        masF = [125,275,525,1025,2075,4175,8375]
        file, _ = QtWidgets.QFileDialog.getOpenFileName(directory='')
        if not file:
            return
        self.ui.lineEdit_2.setText(file)
        try:
            file_np = np.loadtxt(file)
            row = 0
            self.ui.tableWidget.clear()
            #задаем горизонтальные хедеры после clear, так как они обнулятся
            item = QtWidgets.QTableWidgetItem()
            item.setText('F')
            self.ui.tableWidget.setHorizontalHeaderItem(0, item)
            item = QtWidgets.QTableWidgetItem()
            item.setText('Uc_w')
            self.ui.tableWidget.setHorizontalHeaderItem(1, item)
            item = QtWidgets.QTableWidgetItem()
            item.setText('Uw')
            self.ui.tableWidget.setHorizontalHeaderItem(2, item)

            #заполняем таблицу и задаем выравнивание
            for field in masF:
                item = QtWidgets.QTableWidgetItem(str(field))
                item.setTextAlignment(QtCore.Qt.AlignCenter)
                self.ui.tableWidget.setItem(row, 0, item)
                row +=1
            row = 0
            for Pc_w, Pw in file_np:
                column = 1
                item = QtWidgets.QTableWidgetItem(str(Pc_w))
                item.setTextAlignment(QtCore.Qt.AlignCenter)
                self.ui.tableWidget.setItem(row,column,item)
                column = 2
                item = QtWidgets.QTableWidgetItem(str(Pw))
                item.setTextAlignment(QtCore.Qt.AlignCenter)
                self.ui.tableWidget.setItem(row,column,item)
                row +=1
            self.ui.statusbar.showMessage('Файл данных успешно загружен')
        except:
            error = QtWidgets.QMessageBox(QtWidgets.QMessageBox.Information, 'Ошибка',
                                          'Выбран неккоректный файл с данными', QtWidgets.QMessageBox.Ok)
            error.exec()

    #сохранить частоту
    def save_frequence_file(self):
        try:
            item = self.ui.treeWidget.selectedItems()[0]
            item.setCheckState(0, QtCore.Qt.Checked)
            self.ui.treeWidget.update()
            save_folder = str(sys.argv)[2:str(sys.argv).rfind('/') + 1] + 'Files/' + item.text(0) + '.txt'
            arr = np.array([])
            for row in range(0, self.ui.tableWidget.rowCount()):
                for column in range(0, self.ui.tableWidget.columnCount()):
                    item = self.ui.tableWidget.item(row, column)
                    arr = np.append(arr, item.text())
            arr = np.array(arr, dtype=np.float64)
            arr = np.reshape(arr, [self.ui.tableWidget.rowCount(), 3])
            with open(save_folder, 'w') as file:
                for line in arr:
                    for x in line:
                        print(x, file=file, end='\t')
                    print('\n', file=file, end='')
            self.ui.statusbar.showMessage('Данные успешно сохранены в выбранную частоту')
        except:
            error = QtWidgets.QMessageBox(QtWidgets.QMessageBox.Information, 'Ошибка',
                                          'Необходимо выделить частоту ВЧО', QtWidgets.QMessageBox.Ok)
            error.exec()

    #действие при нажатии на частоту в древе
    def item_on_clicked(self, item):
        freq = str(float(item.text(0)))
        self.ui.tableWidget.clear()
        #задаем горизонтальные хедеры после clear, так как они обнулятся
        item = QtWidgets.QTableWidgetItem()
        item.setText('F')
        self.ui.tableWidget.setHorizontalHeaderItem(0, item)
        item = QtWidgets.QTableWidgetItem()
        item.setText('Uc_w')
        self.ui.tableWidget.setHorizontalHeaderItem(1, item)
        item = QtWidgets.QTableWidgetItem()
        item.setText('Uw')
        self.ui.tableWidget.setHorizontalHeaderItem(2, item)
        try:
            open_folder = str(sys.argv)[2:str(sys.argv).rfind('/') + 1] + 'Files/' + freq + '.txt'
            a = np.loadtxt(open_folder)
            for row in range(0, len(a)):
                for column in range(0,len(a[0])):
                    item = QtWidgets.QTableWidgetItem(str(a[row][column]))
                    item.setTextAlignment(QtCore.Qt.AlignCenter)
                    self.ui.tableWidget.setItem(row, column, item)
        except:
            error = QtWidgets.QMessageBox(QtWidgets.QMessageBox.Information, 'Ошибка',
                                          'Открыт пустой файл либо не выбрана частота ВЧО', QtWidgets.QMessageBox.Ok)
            error.exec()

    #сохранить проект
    def save_project(self):
        Files_dir = str(sys.argv)[2:str(sys.argv).rfind('/') + 1] + 'Files/'
        frequences = os.listdir('Files/') #массив файлов проекта
        # если забита хоть одна частота в проекте
        if frequences:
            project_name, _ = QtWidgets.QFileDialog.getSaveFileName(self, 'Сохранить как', '')
            if project_name[-3:] != '.ap':
                project_name = project_name + '.ap'
            if not project_name:
                return
            with open(project_name, 'w') as wfile:
                for freq in frequences:
                    with open(Files_dir + freq, 'r') as rfile:
                        print('['+freq[:-4]+']', file=wfile, end='\n')
                        for line in rfile:
                            print(line, file=wfile,sep='\t', end='')
                        print('\n', file=wfile, end='')
                #запись Pgvc
                print('Pgvc=', str(self.ui.doubleSpinBox.value()), file=wfile, sep='',end='')
            self.ui.statusbar.showMessage('Проект успешно сохранен')
        # если нет частот ошибка
        else:
            error = QtWidgets.QMessageBox(QtWidgets.QMessageBox.Information, 'Ошибка',
                                          'Файл проекта не содержит частот ВЧО', QtWidgets.QMessageBox.Ok)
            error.exec()

    #загрузить проект
    def open_project(self):
        try:
            #Очищаем папку Files
            if os.listdir('Files/'):
                for file in os.listdir('Files/'):
                    os.remove(str(sys.argv)[2:str(sys.argv).rfind('/') + 1] + 'Files/' + file)
            #Рабочие папки
            save_folder = str(sys.argv)[2:str(sys.argv).rfind('/') + 1] + 'Files/'
            open_folder, ok = QtWidgets.QFileDialog.getOpenFileName(self, 'Выберите файл проекта', '', 'AP Files (*.ap)')
            if not ok:
                return
            #массив для частот
            masf = []
            #открываем файл проекта, разделяем на частоты и отдельные файлы для заполнения tableWidget и ListWidget
            with open(open_folder, 'r') as file:
                for line in file:
                    if line[0] == '[':
                        freq = line[1:-2]
                        masf.append(line[1:-2])
                    elif (line[0] != '\n' and line[0] != 'P'):
                        with open(save_folder + freq + '.txt', 'a') as wfile:
                            print(line, file = wfile, end='')
                    elif line[0] == 'P':
                        self.ui.doubleSpinBox.setValue(float(line.strip('Pgvc=')))

        except:
            error = QtWidgets.QMessageBox(QtWidgets.QMessageBox.Information, 'Ошибка',
                                          'Ошибка открытия проекта', QtWidgets.QMessageBox.Ok)
            error.exec()

        #заполняем ListWidget
        try:
            self.ui.treeWidget.clear()
            for x in masf:
                tree_row = QtWidgets.QTreeWidgetItem(self.ui.treeWidget)
                tree_row.setText(0, x)
                tree_row.setTextAlignment(0, QtCore.Qt.AlignCenter)
                tree_row.setFlags(tree_row.flags() | QtCore.Qt.ItemIsEditable)
                #tree_row.setFlags(tree_row.flags()| QtCore.Qt.ItemIsUserCheckable)
                tree_row.setCheckState(0, QtCore.Qt.Checked)
                self.ui.treeWidget.update()
            self.ui.statusbar.showMessage('Проект успешно загружен')
        except:
            error = QtWidgets.QMessageBox(QtWidgets.QMessageBox.Information, 'Ошибка',
                                          'Ошибка заполнения данных в таблицу из загружаемого проекта', QtWidgets.QMessageBox.Ok)
            error.exec()
    #проводить расчет
    def make_rec(self):
        #проверка чекбокса и задано ли значение Pgvc
        if (self.ui.checkBox.isChecked() and self.ui.doubleSpinBox.value() != ''):
            self.ui.statusbar.showMessage('Производится расчет, дождитесь окончания процесса...')
            # try:
            # массив частот ВЧО
            mas_fj = os.listdir('Files/')
            for x in range(0, len(mas_fj)):
                mas_fj[x] = mas_fj[x].rstrip('.txt')

            mas_fj = np.array(mas_fj, dtype=np.float64)
            # print(mas_fj)
            # print(type(mas_fj))

            # переменные ТХ ВЧО противника
            # постоянные переменные из словаря settings_dict
            Pgvc_AR_stac = float(settings_dict['Pgvc_AR_stac'])
            Pgvc_AR_port = float(settings_dict['Pgvc_AR_port'])
            Rprd = float(settings_dict['Rprd'])
            Rprm = float(settings_dict['Rprm'])
            Gym = float(settings_dict['Gym'])
            typeVP = settings_dict['typeVP']

            # значение Pgvc, задается в doublespinbox пользователем
            Pgvc = self.ui.doubleSpinBox.value()

            # калибровочные коэффициенты антенн, из файлов выбранных антенн в папке antens
            # keys = ['PRM_NH', 'PRM_VH', 'PRD_NH', 'PRD_VH']
            # print(antenns_dict['PRM_NH'], antenns_dict['PRM_VH'], antenns_dict['PRD_NH'], antenns_dict['PRD_VH'])
            antenn_folder = str(sys.argv)[2:str(sys.argv).rfind('/') + 1] + 'antens/'
            Ka_prm_NH = np.loadtxt(antenn_folder + antenns_dict['PRM_NH'] + '.txt', delimiter='\t')
            Ka_prm_VH = np.loadtxt(antenn_folder + antenns_dict['PRM_VH'] + '.txt', delimiter='\t')
            Ka_prm_NH = Ka_prm_NH[:-1]  # обрезаем конец, чтобы взять его из антенны ВЧ
            Ka_prm = np.append(Ka_prm_NH, Ka_prm_VH).reshape(len(Ka_prm_NH) + len(Ka_prm_VH), 2)
            Ka_prd_NH = np.loadtxt(antenn_folder + antenns_dict['PRD_NH'] + '.txt', delimiter='\t')
            Ka_prd_VH = np.loadtxt(antenn_folder + antenns_dict['PRD_VH'] + '.txt', delimiter='\t')
            Ka_prd_NH = Ka_prd_NH[:-1]  # обрезаем конец, чтобы взять его из антенны ВЧ
            Ka_prd = np.append(Ka_prd_NH, Ka_prd_VH).reshape(len(Ka_prd_NH) + len(Ka_prd_VH), 2)

            # счетчик f для tablewidget_2
            count_t2 = 0
            # счетчик f для tablewindget_3
            count_t3 = 0
            # self.ui.tableWidget_2.clear()
            # self.ui.tableWidget_3.clear()
            # заполнение tablewidget_2
            # задаем кол-во строк = len(f) * 7
            self.ui.tableWidget_2.setRowCount(len(mas_fj) * 7)
            # rows_t2 = self.ui.tableWidget_2.rowCount()  # строки
            # columns_t2 = self.ui.tableWidget_2.columnCount()  # столбцы
            # заполнение tablewidget_3
            # задаем кол-во строк = len(f) * 2
            self.ui.tableWidget_3.setRowCount(len(mas_fj) * 2)
            # 1 столбец - тип; стац + порт
            item = QtWidgets.QTableWidgetItem('Стационарные')
            item.setTextAlignment(QtCore.Qt.AlignCenter)
            self.ui.tableWidget_3.setItem(0, 0, item)
            item = QtWidgets.QTableWidgetItem('Портативно-\nвозимые')
            item.setTextAlignment(QtCore.Qt.AlignCenter)
            self.ui.tableWidget_3.setItem(0 + len(mas_fj), 0, item)
            rows_t3 = self.ui.tableWidget_3.rowCount()  # строки
            # columns_t3 = self.ui.tableWidget_3.columnCount()  # столбцы

            for fj in mas_fj:
                print(fj)
                Pcw = np.zeros(7)  # 7 - так как известно, что будет 7 чисел
                Pw = np.zeros(7)
                freqs_folder = str(sys.argv)[2:str(sys.argv).rfind('/') + 1] + 'Files/' + str(fj) + '.txt'
                # загрузка данных из файла
                mas = np.loadtxt(freqs_folder)
                for x in range(0, len(mas)):
                    Pcw[x] = mas[x][1]
                    Pw[x] = mas[x][2]

                result = main_copm(fj, Pcw, Pw, Pgvc, Ka_prd, Ka_prm, Pgvc_AR_stac, Pgvc_AR_port, Rprd, Rprm, Gym,
                                   typeVP)
                # print(result)
                # данные с result : R_stac, R_port, Gprd, Gprm, Fi, Pc, qokt
                # заполнение tablewidget_2

                # 1 столбец - частота
                item = QtWidgets.QTableWidgetItem(str(fj))
                item.setTextAlignment(QtCore.Qt.AlignCenter)
                self.ui.tableWidget_2.setItem(count_t2, 0, item)
                # 2 столбец - Gprd
                item = QtWidgets.QTableWidgetItem(str(np.around(result[2], 3)))
                item.setTextAlignment(QtCore.Qt.AlignCenter)
                self.ui.tableWidget_2.setItem(count_t2, 1, item)
                # 3 столбец - Gprm
                item = QtWidgets.QTableWidgetItem(str(np.around(result[3], 3)))
                item.setTextAlignment(QtCore.Qt.AlignCenter)
                self.ui.tableWidget_2.setItem(count_t2, 2, item)
                # 4 столбец - Fi
                for row in range(count_t2, count_t2 + 7):
                    item = QtWidgets.QTableWidgetItem(str(np.around(result[4][row - count_t2], 1)))
                    item.setTextAlignment(QtCore.Qt.AlignCenter)
                    self.ui.tableWidget_2.setItem(row, 3, item)
                # 5 столбец - Pc
                for row in range(count_t2, count_t2 + 7):
                    item = QtWidgets.QTableWidgetItem(str(np.around(result[5][row - count_t2], 3)))
                    item.setTextAlignment(QtCore.Qt.AlignCenter)
                    self.ui.tableWidget_2.setItem(row, 4, item)
                # 6 столбец - qokt
                for row in range(count_t2, count_t2 + 7):
                    item = QtWidgets.QTableWidgetItem(str(np.around(result[6][row - count_t2], 3)))
                    item.setTextAlignment(QtCore.Qt.AlignCenter)
                    self.ui.tableWidget_2.setItem(row, 5, item)
                # объединяем ячейки t2 - задаем спан
                self.ui.tableWidget_2.setSpan(count_t2, 0, 7, 1)
                self.ui.tableWidget_2.setSpan(count_t2, 1, 7, 1)
                self.ui.tableWidget_2.setSpan(count_t2, 2, 7, 1)
                # заполняем tablewidget_3
                # 1 столбец заполнен
                # 2 столбец - частота; стац + порт
                item = QtWidgets.QTableWidgetItem(str(fj))
                item.setTextAlignment(QtCore.Qt.AlignCenter)
                self.ui.tableWidget_3.setItem(count_t3, 1, item)
                item = QtWidgets.QTableWidgetItem(str(fj))
                item.setTextAlignment(QtCore.Qt.AlignCenter)
                self.ui.tableWidget_3.setItem(count_t3 + len(mas_fj), 1, item)
                # 3 столбец - Rar; стац + порт
                item = QtWidgets.QTableWidgetItem(str(result[0]))
                item.setTextAlignment(QtCore.Qt.AlignCenter)
                self.ui.tableWidget_3.setItem(count_t3, 2, item)
                item = QtWidgets.QTableWidgetItem(str(result[1]))
                item.setTextAlignment(QtCore.Qt.AlignCenter)
                self.ui.tableWidget_3.setItem(count_t3 + len(mas_fj), 2, item)
                count_t3 += 1
                count_t2 += 7
            # 4 столбец - maxRar_stac;port
            mas_Rar = np.zeros(len(mas_fj) * 2)
            for row in range(0, rows_t3):
                mas_Rar[row] = self.ui.tableWidget_3.item(row, 2).text()
                # self.ui.tableWidget_3.item
            max_Rar_stac = np.amax(mas_Rar[:len(mas_fj)])
            max_Rar_port = np.amax(mas_Rar[len(mas_fj):])
            item = QtWidgets.QTableWidgetItem(str(max_Rar_stac))
            item.setTextAlignment(QtCore.Qt.AlignCenter)
            self.ui.tableWidget_3.setItem(0, 3, item)
            item = QtWidgets.QTableWidgetItem(str(max_Rar_port))
            item.setTextAlignment(QtCore.Qt.AlignCenter)
            self.ui.tableWidget_3.setItem(len(mas_fj), 3, item)
            # объединяем ячейки t3 - задаем спан
            flag = 0
            for x in range(0, 2):
                self.ui.tableWidget_3.setSpan(flag, 0, len(mas_fj), 1)
                self.ui.tableWidget_3.setSpan(flag, 3, len(mas_fj), 1)
                flag += len(mas_fj)
            # пропорции t3
            header = self.ui.tableWidget_2.horizontalHeader()
            header.setSectionResizeMode(0, QtWidgets.QHeaderView.Stretch)
            # пропорции t3
            header = self.ui.tableWidget_3.horizontalHeader()
            header.setSectionResizeMode(0, QtWidgets.QHeaderView.Stretch)
            self.ui.statusbar.showMessage('Расчет  произведен успешно')

            # except:
            #     error = QtWidgets.QMessageBox(QtWidgets.QMessageBox.Information, 'Ошибка',
            #                               'Ошибка расчета : make_rec',
            #                               QtWidgets.QMessageBox.Ok)
            #     error.exec_()

        # print(main_copm(fj, Pcw, Pw, Pgvc, Kprd, Kprm, Pgvc_AR_stac, Pgvc_AR_port,Rprd,Rprm,Gym,1))
        else:
            error = QtWidgets.QMessageBox(QtWidgets.QMessageBox.Information, 'Ошибка',
                                          'Ошибка расчета : make_rec, проверка не пройдена либо не задано значение Pgvc',
                                          QtWidgets.QMessageBox.Ok)
            error.exec_()

    #экспорт расчета в Word
    def eksport_word_func(self):
        self.ui.statusbar.showMessage('Выполняется экспорт данных в Word...')
        shablons_folder = str(sys.argv)[2:str(sys.argv).rfind('/') + 1] + 'word_shablons/'
        #строки и столбцы таблиц 2 и 3
        rows_t2 = self.ui.tableWidget_2.rowCount()
        column_t2 = self.ui.tableWidget_2.columnCount()
        rows_t3 = self.ui.tableWidget_3.rowCount()
        column_t3 = self.ui.tableWidget_3.columnCount()

        #открываем шаблон документа Word
        #сохраняемый файл
        save_file_name, _ = QtWidgets.QFileDialog.getSaveFileName(self, 'Сохранить как', '')
        if save_file_name[-5:] != '.docx':
            save_file_name = save_file_name + '.docx'
        if not save_file_name:
            return

        try:
            document = Document(shablons_folder + 'VHO_shablon.docx')

            # создаем таблицу подобную tableWidget_2
            table = document.add_table(rows_t2 + 1, column_t2)  # + 1 для хедеров
            table.style = 'Table Grid'
            # table.alignment = WD_TABLE_ALIGNMENT.CENTER #таблица по центру
            table.allow_autofit = True

            # хедеры
            headers_t2 = ('Частота \nтестового \nэлектромагнитного \nсигнала, МГц',
                          'Коэффициент \nусиления \nизлучающей антенны Gпрд(fj), дБ',
                          'Коэффициент \nусиления \nприемной nантенны Gпрм(fj), дБ',
                          'Частота \nтестового \nакустического сигнала Fi, Гц',
                          'Уровень мощности \nмодуляционной \nкомпоненты \nотраженного сигнала \nPc(f, Fi), дБ (мВт)',
                          'Октавная \nЭПР ТС \nσтс (fj, Fi), дБ(м2)')

            # заполняем хедеры таблицы
            for x in range(0, 6):
                cell = table.cell(0, x)
                cell.text = headers_t2[x]

            # 1 столбец
            t2_flag = 0
            for x in range(0, int(rows_t2 / 7)):
                cell = table.cell(t2_flag + 1, 0)
                item = self.ui.tableWidget_2.item(t2_flag, 0)
                cell.text = item.text()
                cell = table.cell(t2_flag + 1, 1)
                item = self.ui.tableWidget_2.item(t2_flag, 1)
                cell.text = item.text()
                cell = table.cell(t2_flag + 1, 2)
                item = self.ui.tableWidget_2.item(t2_flag, 2)
                cell.text = item.text()
                t2_flag += 7

            for column in range(3, column_t2):
                for row in range(0, rows_t2):
                    cell = table.cell(row + 1, column)
                    item = self.ui.tableWidget_2.item(row, column)
                    cell.text = item.text()

            # размер текста 10pt и выравнивание ячеек tableWidget_t2
            for row in table.rows:
                for cell in row.cells:
                    cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    paragraphs = cell.paragraphs
                    for paragraph in paragraphs:
                        for run in paragraph.runs:
                            font = run.font
                            font.size = Pt(10)

            # пустая строка
            document.add_paragraph('')

            # создаем таблицу для tableWidget_3
            table = document.add_table(rows_t3 + 1, column_t3)  # + 1 для хедеров
            table.style = 'Table Grid'
            table.allow_autofit = True

            # хедеры t3
            headers_t3 = ('Вариант размещения \nаппаратуры \nакустической речевой \nразведки',
                          'Частота облучения fj, МГц', 'Дальность разведки Rар-р (fj), м',
                          'Максимальная дальность разведки Rар-р max, м')

            # заполняем хедеры таблицы
            for x in range(0, 4):
                cell = table.cell(0, x)
                cell.text = headers_t3[x]

            # 1 столбец
            # стационарные
            cell = table.cell(1, 0)
            item = self.ui.tableWidget_3.item(0, 0)
            cell.text = item.text()
            # портативно-возимые
            cell = table.cell(int(rows_t3 / 2) + 1, 0)
            item = self.ui.tableWidget_3.item(int(rows_t3 / 2), 0)
            cell.text = item.text()
            # 2 столбец
            for row in range(0, int(rows_t3)):
                cell = table.cell(row + 1, 1)
                item = self.ui.tableWidget_3.item(row, 1)
                cell.text = item.text()
            # 3 столбец
            for row in range(0, int(rows_t3)):
                cell = table.cell(row + 1, 2)
                item = self.ui.tableWidget_3.item(row, 2)
                cell.text = item.text()
            # 4 столбец
            # стационарные
            cell = table.cell(1, 3)
            item = self.ui.tableWidget_3.item(0, 3)
            cell.text = item.text()
            # портативно-возимые
            cell = table.cell(int(rows_t3 / 2) + 1, 3)
            item = self.ui.tableWidget_3.item(int(rows_t3 / 2), 3)
            cell.text = item.text()

            # размер текста 10pt и выравнивание ячеек tableWidget_t3
            for row in table.rows:
                for cell in row.cells:
                    cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    paragraphs = cell.paragraphs
                    for paragraph in paragraphs:
                        for run in paragraph.runs:
                            font = run.font
                            font.size = Pt(10)

            document.save(save_file_name)
            self.ui.statusbar.showMessage('Экспорт данных успешно завершен')
        except:
            error = QtWidgets.QMessageBox(QtWidgets.QMessageBox.Information, 'Ошибка',
                                          'Ошибка формирования расчетного протокола : eksport_word_func',
                                          QtWidgets.QMessageBox.Ok)
            error.exec_()

#меню_настройки_переменных
class menu_options_window(QtWidgets.QDialog, Ui_dialog_options):
    def __init__(self):
        super().__init__()
        self.setupUi(self)

        self.pushButton.clicked.connect(self.restore_settings)

    def restore_settings(self):
        #задаем значения по умолчанию
        self.doubleSpinBox.setValue(float(3))
        self.doubleSpinBox_2.setValue(float(3))
        self.doubleSpinBox_3.setValue(float(0))
        self.doubleSpinBox_4.setValue(float(23))
        self.doubleSpinBox_5.setValue(float(13))

        #combo-box
        self.comboBox.setCurrentIndex(0)  # выбрать 0ый итем

#меню_настройки_антенны
class menu_antenns_options_window(QtWidgets.QDialog, Ui_dialog_antenns):
    def __init__(self):
        super().__init__()
        self.setupUi(self)

        self.pushButton.clicked.connect(self.add_antenna)
        self.pushButton_2.clicked.connect(self.delete_antenna)
        self.listWidget.itemDoubleClicked.connect(self.antena_on_clicked)
        self.pushButton_3.clicked.connect(self.save_antenna)

    #инициализация стандартных настроект антенн, запускается один раз при старте программы
    def init_settings(self):
        global antenn_files
        self.window = menu_antenns_options_window()
        #инициализация данных в listwidget
        settings_folder = str(sys.argv)[2:str(sys.argv).rfind('/') + 1] + 'antens/'
        # print(settings_folder)
        antenn_files = os.listdir(settings_folder)
        #обрезка расширений для отображения
        for x in range(0,len(antenn_files)):
            antenn_files[x] = antenn_files[x].rstrip('.txt')

    def add_antenna(self):
        # ok - флаг
        antena_name, ok = QtWidgets.QInputDialog.getText(QtWidgets.QWidget(),
                                                           'Окно ввода', 'Введите название антенны')
        antena_name = str(antena_name)

        if not ok:
            return

        self.listWidget.addItem(antena_name)
        self.listWidget.update()
        # пополняем данные в combo-boxs
        self.comboBox.addItem(antena_name)
        self.comboBox.update()
        self.comboBox_2.addItem(antena_name)
        self.comboBox_2.update()
        self.comboBox_3.addItem(antena_name)
        self.comboBox_3.update()
        self.comboBox_4.addItem(antena_name)
        self.comboBox_4.update()

        # создаем файл при создании антенны
        with open(str(sys.argv)[2:str(sys.argv).rfind('/') + 1] + 'antens/' + antena_name + '.txt', 'w') as file:
            file.close()
        #обновляем массив с антеннами antenn_files
        antenn_files.append(antena_name)

    def delete_antenna(self):
        items = self.listWidget.selectedItems()
        if not items: return
        for item in items:
            self.listWidget.takeItem(self.listWidget.row(item))
        # также удаляет файл txt
        try:
            os.remove(str(sys.argv)[2:str(sys.argv).rfind('/') + 1] + 'antens/' + item.text() + '.txt')
        except:
            error = QtWidgets.QMessageBox(QtWidgets.QMessageBox.Information, 'Ошибка',
                                          'Ошибка удаления файла настроек антенны delete_antenna,'
                                          'возможно файл уже был удален', QtWidgets.QMessageBox.Ok)
            error.exec()

        #очищаем combo-boxs
        for item in items:
            try:
                index = self.comboBox.findText(item.text())
                self.comboBox.removeItem(index)
                self.comboBox.update()
                index = self.comboBox_2.findText(item.text())
                self.comboBox_2.removeItem(index)
                self.comboBox_2.update()
                index = self.comboBox_3.findText(item.text())
                self.comboBox_3.removeItem(index)
                self.comboBox_3.update()
                index = self.comboBox_4.findText(item.text())
                self.comboBox_4.removeItem(index)
                self.comboBox_4.update()
            except:
                error = QtWidgets.QMessageBox(QtWidgets.QMessageBox.Information, 'Ошибка',
                                              'Ошибка удаления антенны из combo-box delete_antenna',
                                              QtWidgets.QMessageBox.Ok)
                error.exec()
        #удаляем антенну из массива antenn_files
        antenn_files.remove(item.text())

    def antena_on_clicked(self, item):
        antenna = str(item.text())
        self.tableWidget.clear()
        #задаем горизонтальные хедеры после clear, так как они обнулятся
        item = QtWidgets.QTableWidgetItem()
        item.setText('F, МГц')
        self.tableWidget.setHorizontalHeaderItem(0, item)
        item = QtWidgets.QTableWidgetItem()
        item.setText('Ka, дБ')
        self.tableWidget.setHorizontalHeaderItem(1, item)

        try:
            open_folder = str(sys.argv)[2:str(sys.argv).rfind('/') + 1] + 'antens/' + antenna + '.txt'
            a = np.loadtxt(open_folder)
            #если открыт пустой файл, задаем кол-во строк = 20
            if len(a) == 0:
                self.tableWidget.setRowCount(20)
            else:
                self.tableWidget.setRowCount(len(a))
            #задаем значения в таблицу из массива [a]
            for row in range(0, len(a)):
                for column in range(0, len(a[0])):
                    item = QtWidgets.QTableWidgetItem(str(a[row][column]))
                    item.setTextAlignment(QtCore.Qt.AlignCenter)
                    self.tableWidget.setItem(row, column, item)
        except:
            error = QtWidgets.QMessageBox(QtWidgets.QMessageBox.Information, 'Ошибка',
                                          'Открыт пустой файл либо не выбрана антенна', QtWidgets.QMessageBox.Ok)
            error.exec()

    def save_antenna(self):
        try:
            item = self.listWidget.selectedItems()[0]
            save_folder = str(sys.argv)[2:str(sys.argv).rfind('/') + 1] + 'antens/' + item.text() + '.txt'
            arr = np.array([])
            try:
                for row in range(0, self.tableWidget.rowCount()):
                    for column in range(0, self.tableWidget.columnCount()):
                        item = self.tableWidget.item(row, column)
                        arr = np.append(arr, item.text())
                arr = np.array(arr, dtype=np.float64)
            #отлавливаем объект - пустую строку, как окончание цикла
            except:
                rows = int(len(arr)/2)
                arr = np.reshape(arr, [rows, 2])
                with open(save_folder, 'w') as file:
                    for line in arr:
                        for x in line:
                            print(x, file=file, end='\t')
                        print('\n', file=file, end='')
                # print('file created success')
        except:
            error = QtWidgets.QMessageBox(QtWidgets.QMessageBox.Information, 'Ошибка',
                                          'Необходимо выделить антенну', QtWidgets.QMessageBox.Ok)
            error.exec()


def main():
    app = QtWidgets.QApplication(sys.argv)
    window = Main_Window() #главное окно
    window.show()
    window.init_settings() #инициализация настроек переменных
    antenns_window = menu_antenns_options_window()
    antenns_window.init_settings() #инициализация настроек антенн
    #очистка папки Files/ при новом проекте
    if os.listdir('Files/'):
        for file in os.listdir('Files/'):
            os.remove(str(sys.argv)[2:str(sys.argv).rfind('/') + 1] + 'Files/' + file)
    app.exec()

if __name__=='__main__':
    main()